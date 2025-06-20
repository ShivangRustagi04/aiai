import os
import random
import time
import threading
import json
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class CodingInterviewer:
    def __init__(self, model="gemini-2.0-flash"):
        try:
            self.api_key = os.getenv("GEMINI_API_KEY")
            if not self.api_key:
                raise ValueError("Please set the GEMINI_API_KEY in .env file")
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel(model)
            self.conversation_history = []
            self.interview_active = True
            self.coding_questions_asked = 0
            self.max_coding_questions = 2
            self.tech_domains = {
                "frontend": ["React", "Angular", "Vue", "JavaScript", "TypeScript", "CSS", "HTML5"],
                "backend": ["Node.js", "Django", "Spring", "Go", "Rust", "Microservices", "APIs"],
                "AI": ["TensorFlow", "PyTorch", "NLP", "Computer Vision", "LLMs", "Generative AI"],
                "data science": [ "data science","Pandas", "NumPy", "SQL", "Data Visualization", "ETL", "Big Data"],
                "machine learning": ["machine learning","Scikit-learn", "Keras", "Model Deployment", "Feature Engineering"],
                "devops": ["Docker", "Kubernetes", "AWS", "CI/CD", "Terraform", "Monitoring"],
                "mobile": ["Flutter", "React Native", "Swift", "Kotlin", "Mobile UX"],
                "python": ["Python", "Flask", "FastAPI", "Django", "Data Structures", "Algorithms"],
                "java": ["Java", "Spring Boot", "JVM", "Object Oriented Programming", "Collections"],
                "cpp": ["C++", "STL", "Memory Management", "Object Oriented Programming", "Data Structures"]
            }
        except Exception as e:
            print(f"Initialization error: {e}")
            raise

    def _generate_coding_question(self, domain):
        """Generate a coding question based on the candidate's domain."""
        language = domain.capitalize()
        prompt = f"""Generate a medium-level coding problem suitable for a technical interview in {domain}.
        Requirements:
        - Should be solvable in {language}
        - Should take 10-15 minutes to solve
        - Include a clear problem statement
        - Provide input/output examples
        - Should test algorithmic thinking and {domain} knowledge
        Format your response as:
        Problem: [Clear problem statement]
        Example Input: [Sample input]
        Example Output: [Expected output]
        Constraints: [Any constraints or edge cases to consider]
        Generate only the problem, no solution."""
        try:
            response = self.query_gemini(prompt)
            return response.strip() if response else None
        except Exception as e:
            print(f"Error generating coding question: {e}")
            return self._get_fallback_coding_question(domain)

    def _get_fallback_coding_question(self, domain):
        """Fallback coding questions if AI generation fails."""
        fallback_questions = {
            "python": """Problem: Find the two numbers in a list that add up to a target sum.
Example Input: numbers = [2, 7, 11, 15], target = 9
Example Output: [0, 1] (indices of numbers 2 and 7)
Constraints: Each input has exactly one solution, and you may not use the same element twice.""",
            "default": """Problem: Write a function to reverse words in a sentence while keeping the word order.
Example Input: "Hello World Python"
Example Output: "olleH dlroW nohtyP"
Constraints: Preserve spaces between words, handle empty strings gracefully."""
        }
        return fallback_questions.get(domain, fallback_questions["default"])

    def query_gemini(self, prompt):
        """Query the generative AI model."""
        try:
            response = self.model.generate_content(prompt)
            if hasattr(response, 'text'):
                return response.text
            elif hasattr(response, 'candidates') and response.candidates:
                return response.candidates[0].content.parts[0].text
            else:
                return "Could you tell me more about your experience with that?"
        except Exception as e:
            print(f"Gemini API Error: {e}")
            return "Could you elaborate on your experience with that technology?"

    def _save_transcription_to_docx(self, file_path="coding_interview_transcript.docx"):
        """Save the conversation history to a DOCX file."""
        doc = Document()
        doc.add_heading("Coding Interview Transcription", level=1)
        for i, msg in enumerate(self.conversation_history):
            role = msg.get("role", "unknown").capitalize()
            content = msg.get("content", "")
            doc.add_paragraph(f"{role}: {content}")
        doc.save(file_path)
        print(f"Transcript saved to {file_path}")
        return file_path

    def _generate_feedback_from_docx(self, docx_path="coding_interview_transcript.docx"):
        """Generate feedback from the DOCX transcript."""
        try:
            doc = Document(docx_path)
            transcript_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            prompt = f"""
Below is a transcription of a coding interview. Generate concise feedback:
1. Candidate's strengths (max 3 points).
2. Areas for improvement (max 3 points).
3. Overall evaluation (one paragraph).
Transcription:
{transcript_text}
"""
            feedback = self.query_gemini(prompt)
            with open("coding_feedback.json", "w") as f:
                f.write(feedback.strip())
            print("Feedback saved to coding_feedback.json")
        except Exception as e:
            print(f"[ERROR] Feedback generation failed: {e}")

    def _run_interview_logic(self):
        try:
            # Introduction
            print("Hello! Welcome to your coding interview session.")
            print("Let's jump straight into some coding challenges to assess your problem-solving skills.")
            domain = "python"  # Default domain; can be dynamically identified if needed
            while self.coding_questions_asked < self.max_coding_questions and self.interview_active:
                # Generate and present a coding question
                coding_question = self._generate_coding_question(domain)
                print("\nHere's your coding challenge:")
                print(coding_question)
                self.conversation_history.append({"role": "assistant", "content": coding_question})
                print("Describe your approach to solving this problem.")

                # Simulate candidate response (replace with actual listening logic if needed)
                candidate_response = input("\nYour response: ")
                if candidate_response.strip():
                    self.conversation_history.append({"role": "user", "content": candidate_response})
                    feedback = random.choice([
                        "Great explanation!",
                        "Interesting approach!",
                        "That makes sense!"
                    ])
                    print(feedback)
                self.coding_questions_asked += 1

            # Closing remarks
            print("\nThank you for completing the coding challenges!")
            print("Generating feedback and saving the transcription...")
        except Exception as e:
            print(f"Interview error: {e}")
        finally:
            self.interview_active = False
            docx_path = self._save_transcription_to_docx()
            self._generate_feedback_from_docx(docx_path)

    def start_interview(self):
        """Start the coding interview."""
        interview_thread = threading.Thread(target=self._run_interview_logic)
        interview_thread.daemon = True
        interview_thread.start()
        while self.interview_active:
            time.sleep(1)


if __name__ == "__main__":
    try:
        interviewer = CodingInterviewer()
        interviewer.start_interview()
    except Exception as e:
        print(f"Fatal error: {e}")
        print("Please check your environment setup:")
        print("1. GEMINI_API_KEY in .env file")
        print("2. Required Python packages installed")
    print("\nCoding interview session ended.")
    print("Thank you for using the Coding Interview Bot!")