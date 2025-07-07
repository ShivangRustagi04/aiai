import os
import re
import random
import time
from openai import OpenAI
from dotenv import load_dotenv
import cv2
from docx import Document
import pygetwindow as gw
import threading
import subprocess
import tempfile
from datetime import datetime, timedelta
from elevenlabs.client import ElevenLabs
import speech_recognition as sr
import assemblyai as aai

from shared_state import save_to_conversation_history

# Load environment variables
load_dotenv()

class ExpertTechnicalInterviewer:
    def __init__(self, model="gpt-4o-mini-2024-07-18", accent="indian", 
                 client_questions=None, total_duration=80):
        try:
            # Initialize OpenAI client
            self.openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            if not os.getenv("OPENAI_API_KEY"):
                raise ValueError("Please set the OPENAI_API_KEY in .env file")

            self.model = model
            self.interview_state = "introduction"
            self.skill_questions_asked = 0
            self.latest_code_submission = None
            self.last_question = None
            self.just_repeated = False
            self.current_domain = None
            self.conversation_history = []
            self.is_listening = False
            self.interrupted = False
            self.tone_warnings = 0
            self.cheating_warnings = 0
            self.question_count = 0
            self.tab_monitor_ready = False
            self.last_face_detection_time = time.time()
            self.tab_change_detected = False
            self.response_delay = 0.3
            self.accent = accent.lower()
            self.interview_active = True
            self.coding_questions_asked = 0
            self.max_coding_questions = 2
            
            # Initialize ElevenLabs
            elevenlabs_key = os.getenv("ELEVENLABS_API_KEY")
            if not elevenlabs_key:
                raise ValueError("Please set the ELEVENLABS_API_KEY in .env file")
            self.tts_client = ElevenLabs(api_key=elevenlabs_key)
            
            # Initialize speech recognition
            self.recognizer = sr.Recognizer()
            self.recognizer.pause_threshold = 0.8
            self.recognizer.energy_threshold = 4000
            
            # Time management
            self.total_duration = min(max(total_duration, 70), 90)  # Keep between 70-90 minutes
            self.interview_start_time = None
            self.section_start_time = None
            self.current_section = None
            self.client_questions = client_questions or []
            self.used_client_questions = []
            
            # Initialize face detection
            self.face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            self.eye_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_eye.xml')
            
            # Initialize camera
            self.cap = None
            self.camera_active = False
            self.current_coding_question = None
            
            self.tech_domains = {
                "frontend": ["React", "Angular", "Vue", "JavaScript", "TypeScript", "CSS", "HTML5"],
                "backend": ["Node.js", "Django", "Spring", "Go", "Rust", "Microservices", "APIs"],
                "AI": ["TensorFlow", "PyTorch", "NLP", "Computer Vision", "LLMs", "Generative AI"],
                "data science": ["data science", "Pandas", "NumPy", "SQL", "Data Visualization", "ETL", "Big Data"],
                "machine learning": ["machine learning", "Scikit-learn", "Keras", "Model Deployment", "Feature Engineering"],
                "devops": ["Docker", "Kubernetes", "AWS", "CI/CD", "Terraform", "Monitoring"],
                "mobile": ["Flutter", "React Native", "Swift", "Kotlin", "Mobile UX"],
                "python": ["Python", "Flask", "FastAPI", "Django", "Data Structures", "Algorithms"],
                "java": ["Java", "Spring Boot", "JVM", "Object Oriented Programming", "Collections"],
                "cpp": ["C++", "STL", "Memory Management", "Object Oriented Programming", "Data Structures"]
            }
        
            self.non_tech_domains = {
                "edtech": ["Curriculum Design", "Learning Management Systems", "Instructional Design", 
                          "Educational Technology", "Student Engagement", "Assessment Tools"],
                "fintech": ["Digital Payments", "Blockchain", "Risk Management", "Financial Modeling", 
                           "Regulatory Compliance", "Banking Systems"],
                "healthcare": ["Healthcare IT", "Electronic Health Records", "Medical Billing", 
                              "Healthcare Analytics", "Telemedicine", "HIPAA Compliance"]
            }
                
            # Start monitoring threads
            self.monitoring_active = True
            self.last_question = None
            self.tab_monitor_thread = threading.Thread(target=self._monitor_tab_changes)
            self.tab_monitor_thread.daemon = True
            self.tab_monitor_thread.start()

        except Exception as e:
            print(f"Initialization error: {e}")
            raise

    def _check_time_remaining(self, section=None):
        """Check remaining time for current section or total interview"""
        if not self.interview_start_time:
            return float('inf')
            
        elapsed = time.time() - self.interview_start_time
        total_remaining = self.total_duration * 60 - elapsed
        
        if section and self.section_start_time:
            section_elapsed = time.time() - self.section_start_time
            section_duration = self._get_section_duration(section)
            section_remaining = section_duration - section_elapsed
            return min(total_remaining, section_remaining)
            
        return total_remaining

    def _get_section_duration(self, section):
        """Return allocated time for each section in seconds"""
        section_durations = {
            "introduction": 5 * 60,       # 5 minutes
            "background": 7 * 60,         # 7 minutes
            "technical_questions": 30 * 60, # 30 minutes
            "coding_challenge": 20 * 60,    # 20 minutes
            "doubt_clearing": 10 * 60,      # 10 minutes
            "closing": 3 * 60               # 3 minutes
        }
        return section_durations.get(section, 0)

    def _start_section(self, section_name):
        """Mark the start of a new interview section"""
        self.current_section = section_name
        self.section_start_time = time.time()
        print(f"[TIMING] Starting {section_name} section")

    def _adjust_for_time(self, current_phase):
        """Adjust interview flow based on remaining time"""
        remaining = self._check_time_remaining()
        if remaining <= 0:
            self.interview_active = False
            return False
            
        # If we're running short on time, skip to next phase
        if remaining < (self.total_duration * 0.2):  # Last 20% of time
            if current_phase == "technical_questions" and self.coding_questions_asked < self.max_coding_questions:
                return "coding_challenge"
            elif current_phase == "coding_challenge":
                return "doubt_clearing"
                
        return current_phase

    def submit_candidate_code(self, code_string):
        """Save candidate's code submission and ask follow-up questions."""
        self.latest_code_submission = code_string
        if code_string and self.current_coding_question:
            # Check time remaining for coding section
            if self._check_time_remaining("coding_challenge") < 60:  # Less than 1 minute left
                self.speak("We're running short on time, so let's move on after this.", interruptible=False)
                return
                
            # Ask follow-up question about the code
            followup = self._coding_followup(code_string, self._identify_language_from_code(code_string))
            if followup:
                self.speak(followup, interruptible=False)
                time.sleep(0.1)
                answer = self.listen()
                if answer:
                    save_to_conversation_history("user", f"[Follow-up Answer]\n{answer}")

    def _identify_language_from_code(self, code):
        """Simple language detection from code snippet"""
        if "def " in code or "import " in code:
            return "Python"
        elif "class " in code and "{" in code:
            return "Java"
        elif "#include" in code:
            return "C++"
        elif "function " in code or "const " in code:
            return "JavaScript"
        return "Python"  # default

    def wait_after_speaking(self, message, base=0.6, per_word=0.15):
        time.sleep(0.1)  # Always wait 1 second after speaking

    def _give_small_hint(self, question_text):
        hint_prompt = f"""You are an AI coding interviewer. Give a small hint for the following problem.
        It should not reveal the full solution, just nudge the candidate in the right direction.

        Problem:
        {question_text}

        Format: Hint: [short helpful nudge]"""

        hint = self.query_openai(hint_prompt)
        if hint:
            self.speak(hint.strip(), interruptible=False)
    
    def _get_file_extension(self, language):
        return {
            "Python": ".py",
            "Java": ".java",
            "C++": ".cpp",
            "JavaScript": ".js"
        }.get(language, ".txt")

    def _get_fallback_question(self, domain, difficulty):
        """Fallback questions by difficulty level"""
        fallbacks = {
            "easy": {
                "frontend": "Explain how React's virtual DOM works and why it's beneficial.",
                "backend": "What are REST APIs and can you explain the main HTTP methods?",
                "python": "How would you write a function to count word occurrences in a string?",
                "default": "Can you explain the basic architecture of a typical [domain] application?"
            },
            "medium": {
                "frontend": "How would you optimize a React component that's re-rendering too frequently?",
                "backend": "How would you design an API rate limiting system?",
                "python": "How would you implement a memory-efficient data processing pipeline?",
                "default": "What are some common challenges in [domain] and how would you address them?"
            },
            "hard": {
                "frontend": "How would you design a micro-frontend architecture with shared state management?",
                "backend": "How would you design a distributed caching system for a high-traffic API?",
                "python": "How would you implement a thread-safe caching decorator with TTL?",
                "default": "How would you architect a scalable [domain] system for enterprise use?"
            }
        }
        
        return fallbacks.get(difficulty, {}).get(domain, fallbacks[difficulty]["default"])

    def _get_fallback_coding_question(self, domain, difficulty):
        """Fallback coding questions by difficulty level"""
        fallbacks = {
            "easy": {
                "frontend": """Problem: Implement a button counter component in React that increments on click.
                
    Example Input: User clicks button 3 times
    Example Output: Display shows "Count: 3"

    Constraints: Use React hooks""",
                "python": """Problem: Write a function to check if a string is a palindrome.

    Example Input: "racecar"
    Example Output: True

    Constraints: Case insensitive, ignore spaces"""
            },
            "medium": {
                "frontend": """Problem: Create a React form with validation for email and password fields.
                
    Example Input: User enters invalid email
    Example Output: Display error message

    Constraints: Validate email format, password length >= 8""",
                "python": """Problem: Write a function to flatten a nested dictionary.

    Example Input: {'a': 1, 'b': {'c': 2, 'd': 3}}
    Example Output: {'a': 1, 'b.c': 2, 'b.d': 3}

    Constraints: Handle arbitrary nesting levels"""
            },
            "hard": {
                "frontend": """Problem: Implement a debounce hook in React that can be reused across components.

    Example Input: User types quickly in search field
    Example Output: API call made only after 500ms pause

    Constraints: TypeScript, generic implementation""",
                "python": """Problem: Implement a thread-safe LRU cache with TTL expiration.

    Example Input: Cache with size=3, TTL=60s
    Example Output: Evicts least recently used items after size/ttl exceeded

    Constraints: Thread-safe, O(1) operations"""
            }
        }
        
        return fallbacks.get(difficulty, {}).get(domain, """Problem: Write a function to reverse a string.

    Example Input: "hello"
    Example Output: "olleh"

    Constraints: Do not use built-in reverse functions""")

    def _generate_domain_question(self, domain):
        """Generate a domain question based on experience level"""
        if not hasattr(self, 'years_experience'):
            self.years_experience = 0
            
        # Determine difficulty based on experience
        if self.years_experience <= 3:
            difficulty = "easy"
            level_desc = "fundamental"
        elif 4 <= self.years_experience <= 7:
            difficulty = "medium" 
            level_desc = "intermediate"
        else:
            difficulty = "hard"
            level_desc = "advanced"
            
        prompt = f"""Generate one {difficulty} level interview question about {domain} 
        suitable for someone with {self.years_experience} years of experience.
        
        The question should:
        - Test practical, real-world knowledge at {level_desc} level
        - Be clear and concise (1-2 sentences max)  
        - Relate to actual work scenarios
        - Be specific enough to evaluate depth of knowledge
        - Be answerable in 2-3 minutes
        - Must be completely new and not similar to any previously asked questions
        - Should not be a reworded version of any earlier questions
        
        Previously asked questions:
        {[q for q in self.conversation_history if q.get('role') == 'assistant']}
        
        Return only the question, no additional text."""
        
        try:
            question = self.query_openai(prompt)
            return question.strip() if question else None
        except Exception as e:
            print(f"Error generating domain question: {e}")
            return self._get_fallback_question(domain, difficulty)

    def _generate_coding_question(self, domain):
        """Generate coding question based on experience level"""
        if not hasattr(self, 'years_experience'):
            self.years_experience = 0
            
        if self.years_experience <= 3:
            difficulty = "easy"
        elif 4 <= self.years_experience <= 7:
            difficulty = "medium"
        else:
            difficulty = "hard"
            
        prompt = f"""Generate a {difficulty} level coding problem in {domain} 
        suitable for someone with {self.years_experience} years of experience.
        
        Requirements:
        - Should be solvable in 10-15 minutes
        - Include clear problem statement
        - Provide input/output examples
        - Test algorithmic thinking
        - Match {difficulty} difficulty level
        
        Format as:
        Problem: [statement]
        Example Input: [sample] 
        Example Output: [expected]
        Constraints: [any constraints]"""
        
        try:
            response = self.query_openai(prompt)
            return response.strip() if response else None
        except Exception as e:
            print(f"Error generating coding question: {e}")
            return self._get_fallback_coding_question(domain, difficulty)

    def _execute_code(self, language, file_path):
        try:
            if language == "Python":
                result = subprocess.run(["python", file_path], capture_output=True, text=True, timeout=10)
            elif language == "Java":
                compile_result = subprocess.run(["javac", file_path], capture_output=True, text=True)
                if compile_result.returncode != 0:
                    return f"Compile Error:\n{compile_result.stderr}"
                class_name = os.path.basename(file_path).replace(".java", "")
                result = subprocess.run(["java", class_name], capture_output=True, text=True, timeout=10, cwd=os.path.dirname(file_path))
            elif language == "C++":
                exe_path = file_path.replace(".cpp", ".exe")
                compile_result = subprocess.run(["g++", file_path, "-o", exe_path], capture_output=True, text=True)
                if compile_result.returncode != 0:
                    return f"Compile Error:\n{compile_result.stderr}"
                result = subprocess.run([exe_path], capture_output=True, text=True, timeout=10)
            elif language == "JavaScript":
                result = subprocess.run(["node", file_path], capture_output=True, text=True, timeout=10)
            else:
                return "Unsupported language."

            output = ""
            if result.stdout:
                output += f"Output:\n{result.stdout}\n"
            if result.stderr:
                output += f"Errors:\n{result.stderr}\n"
            return output if output else "Code executed successfully (no output)."

        except subprocess.TimeoutExpired:
            return "Error: Code execution timed out (10 seconds limit)"
        except Exception as e:
            return f"Runtime error: {str(e)}"

    def _is_repeat_request(self, text):
        if not text:
            return False
        repeat_phrases = [
            "repeat", "say again", "pardon", "once more",
            "come again", "didn't catch", "hear that"
        ]
        return any(phrase in text.lower() for phrase in repeat_phrases)

    def _run_interview_logic(self):
        try:
            self.interview_start_time = time.time()
            
            # Introduction Section (5 minutes)
            self._start_section("introduction")
            self._conduct_introduction()
            
            # Background Section (7 minutes)
            self._start_section("background")
            self._gather_background()
            
            # Determine interview type
            is_tech_interview = self.current_domain in self.tech_domains
            
            # Technical/Professional Questions Section (30 minutes)
            self._start_section("technical_questions")
            self._ask_client_questions()  # Ask client-provided questions first
            
            # Only proceed with generated questions if we have time
            if not self._should_transition_to_next_section("technical_questions"):
                self._conduct_question_phase(is_tech_interview)
            
            # Coding Challenge Section (20 minutes, tech interviews only)
            if is_tech_interview and self.interview_active:
                self._start_section("coding_challenge")
                self._conduct_coding_challenge()
            
            # Doubt Clearing Section (10 minutes)
            if self.interview_active:
                self._start_section("doubt_clearing")
                self._conduct_doubt_clearing(is_tech_interview)
            
            # Closing Section (3 minutes)
            if self.interview_active:
                self._start_section("closing")
                self._conduct_closing(is_tech_interview)
                
        except Exception as e:
            print(f"Interview error: {e}")
            self.speak("We've encountered a technical issue, but thank you for your participation today!", interruptible=False)
        finally:
            self.interview_active = False
            self.monitoring_active = False
            docx_path = self._save_transcription_to_docx()
            self._generate_feedback_from_docx(docx_path)
            self._stop_camera()

    def _conduct_introduction(self):
        """Handle the introduction section (5 minutes)"""
        self.speak("Hello! I am Gyani. Welcome to your interview session today. I'm excited to chat with you!", interruptible=False)
        msg = "Before we begin, how has your day been so far?"
        self.speak(msg, interruptible=False)
        self.wait_after_speaking(msg)
        day_response = self.listen()

        if day_response:
            save_to_conversation_history("user", day_response)
            self.speak("That's great to hear! I appreciate you taking the time for this session.", interruptible=False)

        msg = "Now, could you please tell me your name and a bit about yourself?"
        self.speak(msg, interruptible=False)
        self.wait_after_speaking(msg)
        introduction = self.listen()

        if introduction:
            self.current_domain = self._identify_tech_domain(introduction)

    def _extract_years_experience(self, text):
        """Extract years of experience from text response"""
        if not text:
            return 0
            
        # Look for patterns like "5 years", "3+ years", "ten years" etc.
        match = re.search(r'(\d+)\s*\+?\s*(?:years|yrs|year)', text)
        if match:
            return int(match.group(1))
            
        # Try to parse written numbers
        number_words = {
            'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
            'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10
        }
        for word, num in number_words.items():
            if word in text.lower():
                return num
                
        return 0  # Default if no experience found

    def _gather_background(self):
        """Gather candidate background information (7 minutes)"""
        is_tech_interview = self.current_domain in self.tech_domains

        # First ask about years of experience
        msg = "Could you tell me how many years of professional experience you have in this field?"
        self.speak(msg, interruptible=False)
        self.wait_after_speaking(msg)
        experience_response = self.listen()
        
        # Extract years of experience from response
        self.years_experience = self._extract_years_experience(experience_response)

        if is_tech_interview:
            msg = "Nice to meet you! Now, I'd love to hear about your technical background and the technologies you enjoy working with."
        else:
            msg = "Nice to meet you! Could you tell me about your professional experience and the domains you've worked in?"
        
        self.speak(msg, interruptible=False)
        self.wait_after_speaking(msg)
        background = self.listen()

        if background:
            self.current_domain = self._identify_tech_domain(background)

    def _should_transition_to_next_section(self, current_section):
        """Determine if we should transition to next section based on time"""
        time_remaining = self._check_time_remaining(current_section)
        min_time_for_next = self._get_section_duration(self._get_next_section(current_section))
        
        # If less than 25% of total time remains or not enough time for next section
        if (time_remaining < (self.total_duration * 60 * 0.25) or 
            time_remaining < min_time_for_next):
            return True
        return False

    def _get_next_section(self, current_section):
        """Get the next section in the interview flow"""
        sections = [
            "introduction",
            "background",
            "technical_questions",
            "coding_challenge",
            "doubt_clearing",
            "closing"
        ]
        try:
            return sections[sections.index(current_section) + 1]
        except (ValueError, IndexError):
            return "closing"

    def _ask_client_questions(self):
        """Ask client-provided questions or generate domain-specific ones"""
        if self.client_questions:
            self.speak("I have some specific questions provided for this interview. Let's begin with those.", interruptible=False)
            time.sleep(0.1)
            
            for question in self.client_questions[:]:
                # Check if we should transition to next section
                if self._should_transition_to_next_section("technical_questions"):
                    return
                    
                if question not in self.used_client_questions:
                    self._ask_question_with_followup(question)
                    self.used_client_questions.append(question)
                    self.client_questions.remove(question)

    def _estimate_experience_level(self):
        """Estimate candidate's experience level based on conversation"""
        # Simple heuristic - count years mentioned
        content = " ".join([msg["content"] for msg in self.conversation_history if msg["role"] == "user"])
        
        if re.search(r'\b([5-9]|\d{2,})\+?\s+years\b', content):
            return "senior"
        elif re.search(r'\b([3-4])\s+years\b', content):
            return "mid"
        return "junior"  # Default

    def _ask_question_with_followup(self, question):
        """Ask a question and follow up based on response"""
        if question == self.last_question and not self.just_repeated:
            return
            
        self.last_question = question
        answer_received = False
        repeat_attempts = 0
        max_repeats = 2

        while not answer_received and repeat_attempts < max_repeats and self.interview_active:
            if not self.just_repeated:
                self.speak(question)
                self.wait_after_speaking(question)
            
            answer = self.listen()
            
            if answer and self._is_repeat_request(answer):
                if repeat_attempts < max_repeats:
                    self.just_repeated = True
                    repeat_attempts += 1
                    rephrased = self._rephrase_question(question)
                    self.speak("Let me rephrase that: " + rephrased)
                    self.last_question = rephrased
                    self.wait_after_speaking(rephrased)
                    continue
                else:
                    placeholder = "[Requested repeat too many times]"
                    save_to_conversation_history("user", placeholder)
                    answer_received = True

            elif not answer or len(answer.split()) <= 3:
                if repeat_attempts < max_repeats - 1:
                    self.speak("Could you please elaborate on that?", interruptible=False)
                else:
                    placeholder = "[Unable to answer after multiple attempts]"
                    save_to_conversation_history("user", placeholder)
                    answer_received = True
            
            elif answer and len(answer.split()) > 4:
                answer_received = True
                
                # Ask follow-up question based on answer
                followup = self._generate_followup_question(question, answer)
                if followup and self._check_time_remaining("technical_questions") > 60:
                    self.speak(followup)
                    self.wait_after_speaking(followup)
                    followup_answer = self.listen()
                    if followup_answer and len(followup_answer.split()) > 4:
                        save_to_conversation_history("assistant", followup)
                
                break

        if answer_received and not self.just_repeated:
            self.question_count += 1  # Increment only for original questions
            save_to_conversation_history("assistant", question)
            
            self.just_repeated = False

    def _generate_followup_question(self, original_question, answer):
            """Generate a relevant follow-up question based on the answer"""
            prompt = f"""Based on this interview exchange, generate relevant follow-up question:
        
        Original Question: {original_question}
        Candidate Answer: {answer}
        
        Requirements:
        - Must be directly related to the answer
        - Should probe deeper into the topic
        - Should be concise (1 sentence)
        - Should be technically/professionally relevant
        - Should help evaluate the candidate's depth of knowledge
        
        Return only the follow-up question."""
            
            followup = self.query_openai(prompt)
            
            if not followup:
                # Fallback follow-up questions
                domain = self.current_domain or "technology"
                followups = {
                    "frontend": "How would you handle edge cases in this scenario?",
                    "backend": "What performance considerations would you make?",
                    "data": "How would this approach scale with larger datasets?",
                    "default": "Can you elaborate on your experience with this?"
                }
                return followups.get(domain, followups["default"])
            
            return followup.strip()

    def _conduct_question_phase(self, is_tech_interview):
        """Conduct the main question phase"""
        question_count = 0
        max_questions = 1  # Adjusted based on time constraints
        
        if is_tech_interview:
            self.speak("Let's start with some technical questions to understand your experience better.", interruptible=False)
        else:
            self.speak("Let's discuss your professional experience in more detail.", interruptible=False)

        while (question_count < max_questions and 
            self.interview_active and 
            not self._should_transition_to_next_section("technical_questions")):
            
            if len(self.conversation_history) > 20:
                self.conversation_history = self.conversation_history[-12:]

            if is_tech_interview:
                system_prompt = f"""As a friendly technical interviewer, ask one engaging question about {self.current_domain or 'technology'} 
                based on this conversation context. The question should:
                - Be completely new and not similar to any previous questions
                - Not be a reworded version of any earlier questions
                - Be one clear question (1 sentence max)
                - Build on what the candidate has already shared
                - Test practical knowledge and experience
                - Be appropriate for their stated experience level
                - Focus on real-world application
                - ask a follow-up about their last answer
                
                Important:
                - DO NOT repeat any questions already asked in this conversation
                - DO NOT ask the same question in different wording
                - ask a follow-up about their last answer
                
                Recent conversation: {' '.join(msg['content'] for msg in self.conversation_history[-3:])}
                
                Previously asked questions:
                {[q for q in self.conversation_history if q.get('role') == 'assistant']}
                
                Generate only the question in a friendly, conversational tone."""
            else:
                system_prompt = f"""As a friendly professional interviewer, ask one engaging question about {self.current_domain or 'professional work'} 
                based on this conversation context. The question should:
                - Be completely new and not similar to any previous questions
                - Not be a reworded version of any earlier questions
                - Be one clear question (1 sentence max)
                - Focus on real-world professional scenarios
                - Test domain knowledge and problem-solving
                - Be appropriate for their stated experience level
                
                Important:
                - DO NOT repeat any questions already asked in this conversation
                - DO NOT ask the same question in different wording
                - If no new questions come to mind, ask a follow-up about their last answer
                
                Recent conversation: {' '.join(msg['content'] for msg in self.conversation_history[-3:])}
                
                Previously asked questions:
                {[q for q in self.conversation_history if q.get('role') == 'assistant']}
                
                Generate only the question in a friendly, conversational tone."""

            response = self.query_openai(system_prompt)
            
            if response:
                msg = response.strip()
                
                if msg == self.last_question and not self.just_repeated:
                    print("[Duplicate] Skipping repeated question.")
                    continue

                self.last_question = msg
                answer_received = False
                repeat_attempts = 0
                max_repeats = 2

                while not answer_received and repeat_attempts < max_repeats:
                    if not self.just_repeated:
                        self.speak(msg)
                        self.wait_after_speaking(msg)
                    
                    answer = self.listen()
                    
                    if answer and self._is_repeat_request(answer):
                        if repeat_attempts < max_repeats:
                            self.just_repeated = True
                            repeat_attempts += 1
                            rephrased = self._rephrase_question(msg)
                            self.speak("Let me rephrase that: " + rephrased)
                            self.last_question = rephrased
                            self.wait_after_speaking(rephrased)
                            continue
                        else:
                            placeholder = "[Requested repeat too many times]"
                            save_to_conversation_history("user", placeholder)
                            answer_received = True

                    elif not answer or len(answer.split()) <= 3:
                        if repeat_attempts < max_repeats - 1:
                            self.speak("Could you please elaborate on that?", interruptible=False)
                        else:
                            placeholder = "[Unable to answer after multiple attempts]"
                            save_to_conversation_history("user", placeholder)
                            answer_received = True
                    
                    elif answer and len(answer.split()) > 4:
                        answer_received = True
                        
                        # Ask follow-up question based on answer
                        followup = self._generate_followup_question(msg, answer)
                        if followup and self._check_time_remaining("technical_questions") > 60:
                            self.speak(followup)
                            self.wait_after_speaking(followup)
                            followup_answer = self.listen()
                            if followup_answer and len(followup_answer.split()) > 4:
                                save_to_conversation_history("assistant", followup)
                        
                        break

                if answer_received and not self.just_repeated:
                    self.question_count += 1  # Increment only for original questions
                    save_to_conversation_history("assistant", msg)
                    
                    self.just_repeated = False

    def _conduct_coding_challenge(self):
        """Conduct coding challenge section with time constraints"""
        self.speak("Great discussion! Now I'd like to give you a couple of coding challenges to see your problem-solving skills in action.", interruptible=False)
        time.sleep(0.1)

        while (self.coding_questions_asked < self.max_coding_questions and 
            self.interview_active and 
            self._check_time_remaining("coding_challenge") > 120):
            
            self.current_coding_question = self._generate_coding_question(self.current_domain or "python")
            save_to_conversation_history("assistant", f"[Coding Challenge Question]\n{self.current_coding_question}")

            self.speak("I've prepared a coding challenge for you. Here's the problem:", interruptible=False)
            time.sleep(0.1)
            print(f"\nCoding Challenge: {self.current_coding_question}")

            self.coding_questions_asked += 1
            hint_offered = False
            start_time = time.time()

            while (self._check_time_remaining("coding_challenge") > 60 and 
                self.interview_active):
                time.sleep(0.1)
                
                # Offer a hint after 2 minutes of inactivity
                if not hint_offered and time.time() - start_time > 120:
                    self.speak("Would you like a small hint to help you get started?", interruptible=False)
                    time.sleep(0.1)
                    response = self.listen()
                    if response and "yes" in response.lower():
                        self._give_small_hint(self.current_coding_question)
                    hint_offered = True

            if not self.interview_active:
                break

            # After code submission, ask follow-up questions
            self.speak("Now let's discuss your solution.", interruptible=False)
            followup = self._coding_followup(self.latest_code_submission, self._identify_language_from_code(self.latest_code_submission))
            if followup:
                self.speak(followup)
                answer = self.listen()
                if answer:
                    # Ask additional follow-up if time permits
                    if self._check_time_remaining("coding_challenge") > 60:
                        second_followup = self._generate_followup_question(followup, answer)
                        if second_followup:
                            self.speak(second_followup)
                            second_answer = self.listen()
            time.sleep(0.1)

    def _coding_followup(self, code, language):
        """Ask follow-up questions about the code submitted by the candidate."""
        prompt = f"""You are an expert software engineer reviewing this {language} code:
        
        ```{code}```
        
        Generate one specific technical follow-up question that:
        1. Tests understanding of the code's logic
        2. Asks about potential improvements
        3. Explores edge cases not handled
        4. Questions algorithmic choices
        5. Is concise (1 sentence)
        
        Example formats:
        - "How would you optimize the time complexity of this solution?"
        - "What edge cases does this code not handle?"
        - "Why did you choose [specific approach] over alternatives?"
        - "How would you modify this to handle [specific scenario]?"
        
        Return only the question."""
        
        try:
            response = self.query_openai(prompt)
            return response.strip() if response else "Can you walk me through your thought process for this solution?"
        except Exception as e:
            print(f"Error generating coding follow-up: {e}")
            return "Can you explain the time complexity of your solution?"

    def _conduct_doubt_clearing(self, is_tech_interview):
        """Conduct doubt clearing session with time constraints"""
        if is_tech_interview:
            self.speak("That was excellent! You've shown great technical knowledge and problem-solving skills.", interruptible=False)
            time.sleep(0.1)

            self.speak("Before we conclude, I'd like to offer you a chance to ask any technical questions you might have.", interruptible=False)
            self.speak("This could be about:", interruptible=False)
            self.speak("1. The coding problems we discussed", interruptible=False)
            self.speak("2. Any of the technical concepts we covered", interruptible=False)
            self.speak("3. Best practices in the field", interruptible=False)
            self.speak("4. Or anything else technical you'd like to discuss", interruptible=False)
        else:
            self.speak("That was excellent! You've shown great professional knowledge and problem-solving skills.", interruptible=False)
            time.sleep(0.1)

            self.speak("Before we conclude, I'd like to offer you a chance to ask any questions you might have about the role or industry.", interruptible=False)
            self.speak("This could be about:", interruptible=False)
            self.speak("1. The professional scenarios we discussed", interruptible=False)
            self.speak("2. Any of the domain concepts we covered", interruptible=False)
            self.speak("3. Industry best practices", interruptible=False)
            self.speak("4. Or anything else you'd like to discuss", interruptible=False)
        
        self.speak("What would you like to ask?", interruptible=False)
        
        questions_asked = 0
        max_questions = 3
        timeout = time.time() + min(self._check_time_remaining("doubt_clearing"), 10 * 60)  # Max 10 minutes

        while (questions_asked < max_questions and 
               self.interview_active and 
               time.time() < timeout):
            question = self.listen()
            if question and len(question.split()) > 3:
                questions_asked += 1
                answer_prompt = f"""Provide a concise but helpful answer to this {'technical' if is_tech_interview else 'professional'} question:
                Question: {question}
                
                Requirements:
                - Keep answer under 4 sentences
                - Be {'technically' if is_tech_interview else 'professionally'} accurate
                - Include one practical example if relevant
                - End by asking if they'd like clarification
                """
                
                answer = self.query_openai(answer_prompt)
                if answer:
                    self.speak(answer, interruptible=False)
                    self.wait_after_speaking(answer)
                    
                    self.speak("Does that answer your question, or would you like me to elaborate?", interruptible=False)
                    followup = self.listen()
                    
                    if followup and "elaborate" in followup.lower():
                        elaboration_prompt = f"""Provide more detailed explanation about:
                        {question}
                        
                        Context:
                        {answer}
                        
                        Requirements:
                        - Go deeper {'technically' if is_tech_interview else 'professionally'}
                        - Include examples
                        - Keep to 5-6 sentences max"""
                        
                        elaboration = self.query_openai(elaboration_prompt)
                        if elaboration:
                            self.speak(elaboration, interruptible=False)
                            self.wait_after_speaking(elaboration)
                
                if questions_asked < max_questions:
                    self.speak("Do you have any other questions?", interruptible=False)
            elif "nothing" in question.lower() or "no questions" in question.lower():
                break

    def _conduct_closing(self, is_tech_interview):
        """Conduct closing remarks"""
        self.speak("Thank you so much for your time today. It was a pleasure talking with you, and I wish you the best of luck!", interruptible=False)

    def _start_camera(self):
        """Start the camera for face detection"""
        if not self.camera_active:
            self.cap = cv2.VideoCapture(0)
            self.camera_active = True

    def _stop_camera(self):
        """Stop the camera"""
        if self.camera_active and self.cap:
            self.cap.release()
            self.camera_active = False

    def _restart_camera(self):
        try:
            if self.cap:
                self.cap.release()
            self.cap = cv2.VideoCapture(0)
            self.cap.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
            self.cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)
            self.camera_active = True
        except Exception as e:
            print(f"Camera restart failed: {e}")

    def _monitor_tab_changes(self):
        while not self.tab_monitor_ready:
            time.sleep(0.5)
        
        try:
            initial_window = gw.getActiveWindow()
            initial_title = initial_window.title if initial_window else "Interview Window"
        except:
            initial_window = None
            initial_title = "Interview Window"
        
        warning_given = False
        
        while self.monitoring_active and self.interview_active:
            try:
                current_window = gw.getActiveWindow()
                current_title = current_window.title if current_window else initial_title
                
                # Only trigger if:
                # 1. We have a valid window
                # 2. The title has actually changed significantly
                # 3. The new title doesn't contain system/keywords
                if (current_window and initial_window and 
                    current_title != initial_title and
                    not any(x in current_title.lower() for x in ["notification", "system", "settings"])):
                    
                    if not warning_given:  # Only warn once per change
                        self.tab_change_detected = True
                        self._handle_cheating_attempt("tab_change")
                        warning_given = True
                else:
                    warning_given = False
                    
                time.sleep(3)  # Longer delay between checks
                
            except Exception as e:
                print(f"Window monitoring error: {e}")
                time.sleep(3)

    def _handle_cheating_attempt(self, cheat_type):
        """Handle different types of cheating attempts"""
        self.cheating_warnings += 1
        
        if self.cheating_warnings >= 3:
            self.speak("Multiple concerning behaviors detected. The interview will now conclude.", interruptible=False)
            self.interview_active = False
            return
            
        responses = {
            "tab_change": "Please stay focused on the interview window and avoid switching to other applications."
        }
        
        if cheat_type in responses:
            self.speak(f"Gentle reminder: {responses[cheat_type]} This is notice {self.cheating_warnings} of 3.", interruptible=False)

    def __del__(self):
        """Clean up resources"""
        self.interview_active = False
        self.monitoring_active = False
        self._stop_camera()
        if hasattr(self, 'face_monitor_thread'):
            self.face_monitor_thread.join(timeout=1)
        if hasattr(self, 'tab_monitor_thread'):
            self.tab_monitor_thread.join(timeout=1)

    def speak(self, text, interruptible=True):
        if self.interrupted:
            self.interrupted = False
            return

        print(f"Interviewer: {text}")
        save_to_conversation_history("assistant", text)

        try:
            # Generate audio with ElevenLabs
            audio = self.tts_client.text_to_speech.stream(
                voice_id="xnx6sPTtvU635ocDt2j7",
                optimize_streaming_latency="0",
                output_format="mp3_44100_128",
                text=text,
                model_id="eleven_turbo_v2"
            )
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp_file:
                for chunk in audio:
                    if chunk:
                        tmp_file.write(chunk)
                tmp_file_path = tmp_file.name
            # Play using FFmpeg
            subprocess.run([
                "ffmpeg",
                "-hide_banner",
                "-loglevel", "error",
                "-i", tmp_file_path,
                "-f", "wav",
                "pipe:1"
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            # Play audio directly
            subprocess.run([
                "ffplay",
                "-nodisp",
                "-autoexit",
                "-loglevel", "quiet",
                tmp_file_path
            ])
            
            # Clean up
            os.unlink(tmp_file_path)
            
            # Calculate wait time
            word_count = len(text.split())
            wait_time = max(0.5, word_count * 0.2)
            time.sleep(wait_time)

        except Exception as e:
            print(f"Audio playback error: {e}")
            print("[TTS Failed] Audio could not be played")

    def listen(self, max_attempts=3):
        """Listen for user response using microphone with faster STT options"""
        for attempt in range(max_attempts):
            try:
                print("\nListening... (Speak now)")
                
                with sr.Microphone() as source:
                    self.recognizer.adjust_for_ambient_noise(source)
                    audio = self.recognizer.listen(source, timeout=10, phrase_time_limit=30)
                
                print("Processing your speech...")
                
                # Try Google Web Speech API first (fastest but requires internet)
                try:
                    text = self.recognizer.recognize_google(audio)
                    print(f"Candidate: {text}")
                    
                    # Process tone detection
                    tone = self._detect_tone(text)
                    if tone != "professional":
                        self.handle_improper_tone(tone)
                        placeholder = "[Response had non-professional tone]"
                        save_to_conversation_history("user", placeholder)
                        return placeholder
                    
                    save_to_conversation_history("user", text)
                    return text
                    
                except sr.UnknownValueError:
                    print("Google Web Speech could not understand audio")
                    raise             
            except Exception as e:
                print(f"Speech recognition error: {e}")
                if attempt < max_attempts - 1:
                    self.speak("There was an issue. Please try again.", interruptible=False)
                    time.sleep(2)
        
        placeholder = "[Response unclear after multiple attempts]"
        save_to_conversation_history("user", placeholder)
        self.speak("Let's continue with the next part.", interruptible=False)
        return placeholder

    def _rephrase_question(self, question):
        """Rephrase the given question while keeping the same meaning"""
        prompt = f"""Rephrase this interview question to make it clearer while keeping the same meaning:
        Original: {question}
        
        Requirements:
        - Keep technical accuracy
        - Maintain same difficulty level
        - Don't change the core concept being tested
        - Make it slightly different wording
        - Keep it one sentence
        
        Return only the rephrased question."""
        
        rephrased = self.query_openai(prompt)
        return rephrased.strip() if rephrased else question

    def _detect_tone(self, text):
        if not text:
            return "professional"
            
        text_lower = re.sub(r'\s+', ' ', text.lower().strip())
        
        arrogant_keywords = [
            r'\bobviously\b', r'\beveryone knows\b', r'\bchild\'?s play\b',
            r'\bthat\'?s easy\b', r'\btrivial\b', r'\bwaste of time\b'
        ]
        
        rude_patterns = [
            r'\byou don\'?t understand\b', r'\bthat\'?s stupid\b', r'\bdumb question\b',
            r'\bare you serious\b', r'\bthis is ridiculous\b', r'\bwho cares\b'
        ]
        
        for pattern in arrogant_keywords:
            if re.search(pattern, text_lower):
                return "arrogant"
                
        for pattern in rude_patterns:
            if re.search(pattern, text_lower):
                return "rude"
                
        return "professional"

    def handle_improper_tone(self, tone):
        self.tone_warnings += 1
        
        if self.tone_warnings >= 2:
            self.speak("I appreciate your participation, but let's maintain a professional tone throughout our conversation.", interruptible=False)
            return
            
        responses = {
            "arrogant": [
                "I appreciate your confidence! Let's channel that into demonstrating your technical knowledge.",
                "Great confidence! Now let's see how you apply that expertise to solve problems.",
            ],
            "rude": [
                "I understand technical interviews can be stressful. Let's take a moment and continue professionally.",
                "No worries, let's refocus on showcasing your technical abilities.",
            ]
        }
        
        if tone in responses:
            response = random.choice(responses[tone])
            self.speak(response, interruptible=False)
            time.sleep(0.1)

    def query_openai(self, prompt):
        try:
            if "generate one engaging question" in prompt.lower():
                prompt += "\n\nImportant: Do not repeat any questions already asked in this conversation."
                
            response = self.openai_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert technical interviewer.Ask engaging questions. Do not repeat questions already asked . Do not ask same question in differnt way unless candidate asks to repeat the question.Should not be a reworded version of any earlier questions.next question should be related to the candidate's last answer. continue for whole technical discussion phase"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=500
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"OpenAI API Error: {e}")
            return "Could you elaborate on your experience with that technology?"

    def _identify_tech_domain(self, text):
        if not text:
            return None

        text_lower = text.lower()
        domain_matches = {
            "data science": "data science",
            "machine learning": "machine learning",
            "data scientist": "data science",
            "ml engineer": "machine learning"
        }
        
        for phrase, domain in domain_matches.items():
            if phrase in text_lower:
                print(f"[Domain Detection] Direct match found: {domain}")
                return domain

        # Then proceed with the regular scoring system
        domain_scores = {domain: 0 for domain in self.tech_domains}
        
        for domain, keywords in self.tech_domains.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    domain_scores[domain] += 2
                elif re.search(r'\b' + re.escape(keyword.lower()) + r'\b', text_lower):
                    domain_scores[domain] += 1

        best_tech_domain, tech_score = max(domain_scores.items(), key=lambda x: x[1])

        if tech_score >= 1:  # Lowered threshold from 2 to 1
            print(f"[Domain Detection] Detected technical domain: {best_tech_domain}")
            return best_tech_domain

        # Check non-technical domains
        domain_scores = {domain: 0 for domain in self.non_tech_domains}
        for domain, keywords in self.non_tech_domains.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    domain_scores[domain] += 2
                elif re.search(r'\b' + re.escape(keyword.lower()) + r'\b', text_lower):
                    domain_scores[domain] += 1

        best_non_tech_domain, non_tech_score = max(domain_scores.items(), key=lambda x: x[1])

        if non_tech_score >= 1:  # Lowered threshold from 2 to 1
            print(f"[Domain Detection] Detected non-technical domain: {best_non_tech_domain}")
            return best_non_tech_domain

        print("[Domain Detection] No domain match found.")
        return None

    def _save_transcription_to_docx(self, file_path="interview_transcript.docx"):
        doc = Document()
        doc.add_heading("Interview Transcription", level=1)
        
        for i, msg in enumerate(self.conversation_history):
            role = msg.get("role", "unknown").capitalize()
            content = msg.get("content", "")
            doc.add_paragraph(f"{role}: {content}")
        
        doc.save(file_path)
        print(f"Transcript saved to {file_path}")
        return file_path

    def _generate_feedback_from_docx(self, docx_path="interview_transcript.docx"):
        try:
            doc = Document(docx_path)
            transcript_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            prompt = f"""
                        Below is a transcription of an interview. Perform the following tasks STRICTLY as described:
            1. Extract the interviewer's questions and the candidate's answers:
            - Ignore filler words like "okay," "hmm," "uh," etc., unless part of a meaningful question/answer.
            - Only include complete sentences or meaningful phrases.
            - Skip any exchange where:
                - The question is filler (e.g., "Okay", "Hmm", "Got it").
                - The answer is too short, incomplete, or irrelevant (e.g., "Yes", "No", "Maybe", "I think so").
            2. Categorize each question under a generalized skill category (e.g., Python, AI, JavaScript, Machine Learning, etc.).
            3. For each skill category:
            - Summarize the candidate's performance concisely (STRICTLY within 900 characters). If exceeded, truncate and add "..." at the end.
            4. For each extracted question-answer pair:
            - Include start and end timestamps in seconds (relative to interview start).
            - Group questions by skill area into a single block.
            - Ensure "que" (question) is within 900 characters. If exceeded, truncate and add "...".
            - Ensure "ans" (answer) is within 4000 characters. If exceeded, truncate and add "...".
            - Also provide a concise alternate ideal answer (under 1000 characters) if the answer is weak.
            5. Provide an overall evaluation:
            - Candidate strengths (STRICTLY within 400 characters). If exceeded, truncate and add "...".
            - Points of improvement (STRICTLY within 400 characters). If exceeded, truncate and add "..."
            6. Rate the candidate on:
            - Communication: Choose one — poor, average, good, excellent.
            - Attitude: Choose one — poor, average, good, excellent.
            7. If the candidate asked to repeat a question or said they didn't understand:
            - Indicate that and include the **rephrased question** that was asked instead.
            - If they still could not answer, generate a **model answer** that could have helped.
            - Mark such entries under `"alternate_explanation": "..."` in the output.
            8. Create a feedback form titled `"feedback_form"` containing:
            - `"summary"`: a concise paragraph (300–500 characters) summarizing the overall performance
            - `"key_areas_to_improve"`: list of 3 suggestions
            - `"recommended_study_links"`: list of URLs or topics (max 5)
            ...
            Transcription:
            {transcript_text}
            """

            feedback = self.query_openai(prompt)
            
            with open("final_interview_feedback.json", "w") as f:
                f.write(feedback.strip())
            print("Feedback saved to final_interview_feedback.json")
        
        except Exception as e:
            print(f"[ERROR] Feedback generation failed: {e}")

    def start_interview(self):
        # Start tab monitoring
        def enable_tab_monitor():
            time.sleep(3)
            self.tab_monitor_ready = True

        threading.Thread(target=enable_tab_monitor, daemon=True).start()

        # Start the interview logic
        interview_thread = threading.Thread(target=self._run_interview_logic)
        interview_thread.daemon = True
        interview_thread.start()

        # Keep the main thread alive while interview is active
        while self.interview_active:
            time.sleep(0.1)

if __name__ == "__main__":
    try:
        # Example usage with client-provided questions
        client_questions = []
        
        interviewer = ExpertTechnicalInterviewer(
            total_duration=80  # 80 minute interview
        )
        interviewer.start_interview()
    except Exception as e:
        print(f"Fatal error: {e}")
        print("Please check your environment setup:")
        print("1. OPENAI_API_KEY in .env file")
        print("2. ELEVENLABS_API_KEY in .env file")
        print("3. Microphone and camera permissions")
        print("4. Required Python packages installed")
    
    print("\nInterview session ended. Camera automatically turned off.")
    print("Thank you for using the Enhanced Technical Interview Bot!")