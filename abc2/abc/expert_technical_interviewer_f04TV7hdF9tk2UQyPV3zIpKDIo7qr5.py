import os
import re
import random
import time
import requests
import speech_recognition as sr
from gtts import gTTS
import pygame
import queue
from io import BytesIO
import google.generativeai as genai
from dotenv import load_dotenv
import cv2
import numpy as np
from docx import Document
import pygetwindow as gw
import threading
import wave
from scipy.io import wavfile
import subprocess
import tempfile
import sys
import boto3
from sentence_transformers import SentenceTransformer
import faiss  # For vector similarity search
import json
from datetime import datetime, timedelta

# Load environment variables
load_dotenv()

class ExpertTechnicalInterviewer:
    def __init__(self, model="gemini-2.0-flash", accent="indian"):
        try:
            self.api_key = os.getenv("GEMINI_API_KEY")
            if not self.api_key:
                raise ValueError("Please set the GEMINI_API_KEY in .env file")
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel(model)
            self.interview_state = "introduction"
            self.skill_questions_asked = 0
            self.latest_code_submission = None
            self.last_question = None
            self.just_repeated = False
            self.current_domain = None
            self.conversation_history = []
            self.recognizer = sr.Recognizer()
            self.microphone = sr.Microphone()
            self.is_listening = False
            self.interrupted = False
            self.recognizer.pause_threshold = 0.6
            self.recognizer.phrase_threshold = 0.2
            self.tone_warnings = 0
            self.cheating_warnings = 0
            self.filler_phrases = []
            self.tab_monitor_ready = False
            self.last_face_detection_time = time.time()
            self.tab_change_detected = False
            self.response_delay = 0.3
            self.accent = accent.lower()
            self.interview_active = True
            self.coding_questions_asked = 0
            self.max_coding_questions = 2
            self.polly = boto3.client(
                "polly",
                region_name=os.getenv("AWS_REGION"),
                aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY")
            )
            # Initialize face detection
            self.face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            self.eye_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_eye.xml')
            # Initialize camera
            self.cap = None
            self.camera_active = False
            self.current_coding_question = None
            self.tech_domains = {
                "frontend": ["React", "Angular", "Vue", "JavaScript", "TypeScript", "CSS", "HTML5"],
                "backend": ["Node.js", "Django", "Spring", "Go", "Rust", "Microservices", "APIs"],
                "AI": ["TensorFlow", "PyTorch", "NLP", "Computer Vision", "LLMs", "Generative AI"],
                "data science": ["data science", "Pandas", "NumPy", "SQL", "Data Visualization", "ETL", "Big Data"],
                "machine learning": ["machine learning", "Scikit-learn", "Keras", "Model Deployment", "Feature Engineering"],
                "devops": ["Docker", "Kubernetes", "AWS", "CI/CD", "Terraform", "Monitoring"],
                "mobile": ["Flutter", "React Native", "Swift", "Kotlin", "Mobile UX"],
                "python": ["Python", "Flask", "FastAPI", "Data Structures", "Algorithms"],
                "java": ["Java", "Spring Boot", "JVM", "Object Oriented Programming", "Collections"],
                "cpp": ["C++", "STL", "Memory Management", "Object Oriented Programming", "Data Structures"]
            }
            self.non_tech_domains = {
                "edtech": ["Curriculum Design", "Learning Management Systems", "Instructional Design", 
                          "Educational Technology", "Student Engagement", "Assessment Tools"],
                "fintech": ["Digital Payments", "Blockchain", "Risk Management", "Financial Modeling", 
                           "Regulatory Compliance", "Banking Systems"],
                "healthcare": ["Healthcare IT", "Electronic Health Records", "Medical Billing", 
                              "Healthcare Analytics", "Telemedicine", "HIPAA Compliance"],
                "banking": ["Retail Banking", "Investment Banking", "Wealth Management", 
                           "Loan Processing", "Anti-Money Laundering", "Financial Analysis"],
                "insurance": ["Underwriting", "Claims Processing", "Actuarial Science", 
                             "Risk Assessment", "Policy Administration", "Customer Service"]
            }

            # Initialize PyGame mixer with error handling
            try:
                pygame.mixer.init(frequency=22050, size=-16, channels=2, buffer=512)
            except pygame.error as e:
                print(f"PyGame mixer initialization failed: {e}")
                raise RuntimeError("Audio system initialization failed")

            # Start monitoring threads
            self.monitoring_active = True
            self.last_question = None
            self.face_monitor_thread = threading.Thread(target=self._monitor_face_and_attention)
            self.face_monitor_thread.daemon = True
            self.face_monitor_thread.start()
            self.tab_monitor_thread = threading.Thread(target=self._monitor_tab_changes)
            self.tab_monitor_thread.daemon = True
            self.tab_monitor_thread.start()

            # Interview Timing Configuration
            self.time_limits = {
                "introduction": 5,
                "technical_discussion": 30,
                "coding_challenge": 25,
                "doubt_clearing": 10,
                "conclusion": 5
            }
            self.total_time_limit = sum(self.time_limits.values())  # ~75 min
            self.interview_start_time = None
            self.client_questions = []

        except Exception as e:
            print(f"Initialization error: {e}")
            raise

    def _check_time_remaining(self, section):
        """Check if there's time remaining for the current section"""
        if self.interview_start_time is None:
            return True
        time_elapsed = (datetime.now() - self.interview_start_time).total_seconds() / 60
        time_used_in_section = time_elapsed - sum(self.time_limits[section] 
                                                for i, section in enumerate(self.time_limits) 
                                                if i < list(self.time_limits.keys()).index(section))
        return time_used_in_section < self.time_limits[section]

    def load_client_questions(self, file_path):
        """Load client-provided questions from a JSON file"""
        try:
            with open(file_path, 'r') as f:
                self.client_questions = json.load(f)
            print(f"Loaded {len(self.client_questions)} client-provided questions")
            return True
        except Exception as e:
            print(f"Error loading client questions: {e}")
            return False

    def set_custom_time_limits(self, time_config):
        """Set custom time limits for each interview section"""
        for section, time in time_config.items():
            if section in self.time_limits:
                self.time_limits[section] = time
        self.total_time_limit = sum(time_config.values()) + 5  # Add buffer time

    def submit_candidate_code(self, code_string):
        """Save candidate's code submission and ask follow-up questions."""
        self.latest_code_submission = code_string
        if code_string and self.current_coding_question:
            # Ask follow-up question about the code
            followup = self._coding_followup(code_string, self._identify_language_from_code(code_string))
            if followup:
                self.speak(followup, interruptible=False)
                time.sleep(1)  # 1 second pause
                answer = self.listen()
                if answer:
                    self.conversation_history.append({
                        "role": "user", 
                        "content": f"[Follow-up Answer]\n{answer}"
                    })

    def _identify_language_from_code(self, code):
        """Simple language detection from code snippet"""
        if "def " in code or "import " in code:
            return "Python"
        elif "class " in code and "{" in code:
            return "Java"
        elif "#include" in code:
            return "C++"
        elif "function " in code or "const " in code:
            return "JavaScript"
        return "Python"  # default

    def wait_after_speaking(self, message, base=0.6, per_word=0.15):
        time.sleep(1)  # Always wait 1 second after speaking

    def _give_small_hint(self, question_text):
        hint_prompt = f"""You are an AI coding interviewer. Give a small hint for the following problem.
        It should not reveal the full solution, just nudge the candidate in the right direction.
        Problem:
        {question_text}
        Format: Hint: [short helpful nudge]"""
        hint = self.query_gemini(hint_prompt)
        if hint:
            self.speak(hint.strip(), interruptible=False)

    def _get_file_extension(self, language):
        return {
            "Python": ".py",
            "Java": ".java",
            "C++": ".cpp",
            "JavaScript": ".js"
        }.get(language, ".txt")

    def _generate_non_tech_question(self, domain):
        """Generate non-technical questions for professional interviews"""
        prompt = f"""Generate one professional interview question about {domain} that:
        - Tests domain knowledge
        - Is relevant to real work situations
        - Is clear and concise
        - Is appropriate for an interview setting
        Generate only the question, no additional text."""
        try:
            response = self.query_gemini(prompt)
            return response.strip() if response else None
        except Exception as e:
            print(f"Error generating non-tech question: {e}")
            return f"Can you describe your experience working in {domain}?"

    def _execute_code(self, language, file_path):
        try:
            if language == "Python":
                result = subprocess.run(["python", file_path], capture_output=True, text=True, timeout=10)
            elif language == "Java":
                compile_result = subprocess.run(["javac", file_path], capture_output=True, text=True)
                if compile_result.returncode != 0:
                    return f"Compile Error:\n{compile_result.stderr}"
                class_name = os.path.basename(file_path).replace(".java", "")
                result = subprocess.run(["java", class_name], capture_output=True, text=True, timeout=10, cwd=os.path.dirname(file_path))
            elif language == "C++":
                exe_path = file_path.replace(".cpp", ".exe")
                compile_result = subprocess.run(["g++", file_path, "-o", exe_path], capture_output=True, text=True)
                if compile_result.returncode != 0:
                    return f"Compile Error:\n{compile_result.stderr}"
                result = subprocess.run([exe_path], capture_output=True, text=True, timeout=10)
            elif language == "JavaScript":
                result = subprocess.run(["node", file_path], capture_output=True, text=True, timeout=10)
            else:
                return "Unsupported language."
            output = ""
            if result.stdout:
                output += f"Output:\n{result.stdout}\n"
            if result.stderr:
                output += f"Errors:\n{result.stderr}\n"
            return output if output else "Code executed successfully (no output)."
        except subprocess.TimeoutExpired:
            return "Error: Code execution timed out (10 seconds limit)"
        except Exception as e:
            return f"Runtime error: {str(e)}"

    def _is_repeat_request(self, text):
        if not text:
            return False
        repeat_phrases = [
            "repeat", "say again", "pardon", "once more",
            "come again", "didn't catch", "hear that"
        ]
        return any(phrase in text.lower() for phrase in repeat_phrases)

    def _run_interview_logic(self):
        try:
            self.interview_start_time = datetime.now()

            # Friendly introduction within time limit
            if self._check_time_remaining("introduction"):
                self.speak("Hello! I am Gyani. Welcome to your interview session today. I'm excited to chat with you!", interruptible=False)
                time.sleep(1)
                msg = "Before we begin, how has your day been so far?"
                self.speak(msg, interruptible=False)
                self.wait_after_speaking(msg)
                day_response = self.listen()
                if day_response:
                    self.conversation_history.append({"role": "user", "content": day_response})
                    self.speak("That's great to hear! I appreciate you taking the time for this session.", interruptible=False)

            # Technical discussion phase
            if self._check_time_remaining("technical_discussion"):
                msg = "Now, could you please tell me your name and a bit about yourself?"
                self.speak(msg, interruptible=False)
                self.wait_after_speaking(msg)
                introduction = self.listen()
                
                if introduction:
                    self.conversation_history.append({"role": "user", "content": introduction})
                    self.current_domain = self._identify_tech_domain(introduction)
                    
                    # Use client-provided questions if available
                    questions_to_ask = self.client_questions if self.client_questions else []
                    
                    # Add generated questions if needed
                    while len(questions_to_ask) < 5:  # Ensure minimum of 5 questions
                        new_question = self._generate_coding_question(self.current_domain)
                        if new_question and new_question not in questions_to_ask:
                            questions_to_ask.append(new_question)
                            
                    # Ask questions with time constraints
                    for question in questions_to_ask:
                        if not self._check_time_remaining("technical_discussion"):
                            break
                            
                        self.speak(question, interruptible=False)
                        answer = self.listen()
                        
                        if answer:
                            self.conversation_history.append({"role": "user", "content": answer})
                            
                            # Generate follow-up question
                            followup = self._generate_followup_question(question, answer)
                            if followup:
                                self.speak(followup, interruptible=False)
                                followup_answer = self.listen()
                                
                                if followup_answer:
                                    self.conversation_history.append({"role": "user", "content": followup_answer})

            # Coding challenge phase with time constraint
            if self._check_time_remaining("coding_challenge") and self.coding_questions_asked < self.max_coding_questions:
                self._conduct_coding_challenge()
                
            # Doubt clearing session with time constraint
            if self._check_time_remaining("doubt_clearing"):
                self.speak("Before we conclude, do you have any questions?", interruptible=False)
                response = self.listen()
                if response:
                    self.speak("Let me address that for you.", interruptible=False)

            # Conclusion within time limit
            if self._check_time_remaining("conclusion"):
                self.speak("Thank you for your time. The interview is now concluding.", interruptible=False)

        except Exception as e:
            print(f"Interview error: {e}")
            self.speak("We've encountered a technical issue, but thank you for your participation today!", interruptible=False)

    def _generate_followup_question(self, original_question, answer):
        """Generate a follow-up question based on the answer"""
        prompt = f"""Given this interview question and answer, generate one follow-up question that:
        - Tests deeper understanding of the same concept
        - Is specific to the answer provided
        - Helps assess the candidate's depth of knowledge
        - Is clear and concise
        
        Question: {original_question}
        Answer: {answer}
        
        Generate only the follow-up question."""
        
        return self.query_gemini(prompt)

    def _conduct_coding_challenge(self):
        self.speak("Great discussion! Now I'd like to give you a couple of coding challenges to see your problem-solving skills in action.", interruptible=False)
        time.sleep(1)  # 1 second pause
        while self.coding_questions_asked < self.max_coding_questions and self.interview_active:
            self.current_coding_question = self._generate_coding_question(self.current_domain or "python")
            self.conversation_history.append({
                "role": "assistant",
                "content": f"[Coding Challenge Question]\n{self.current_coding_question}"
            })
            self.speak("I've prepared a coding challenge for you. Here's the problem:", interruptible=False)
            time.sleep(1)  # 1 second pause after speaking question
            print(f"\nCoding Challenge: {self.current_coding_question}")
            self.coding_questions_asked += 1
            hint_offered = False
            start_time = time.time()
            while self.coding_questions_asked < self.max_coding_questions and self.interview_active:
                time.sleep(1)
                # Offer a hint after 2 minutes of inactivity
                if not hint_offered and time.time() - start_time > 120:
                    self.speak("Would you like a small hint to help you get started?", interruptible=False)
                    time.sleep(1)  # 1 second pause
                    response = self.listen()
                    if response and "yes" in response.lower():
                        self._give_small_hint(self.current_coding_question)
                    hint_offered = True
            if not self.interview_active:
                break
            time.sleep(1)

    def _start_camera(self):
        """Start the camera for face detection"""
        if not self.camera_active:
            self.cap = cv2.VideoCapture(0)
            self.camera_active = True

    def _stop_camera(self):
        """Stop the camera"""
        if self.camera_active and self.cap:
            self.cap.release()
            self.camera_active = False

    def _monitor_face_and_attention(self):    
        multiple_faces_warning_given = False
        looking_away_warning_given = False
        while self.monitoring_active and self.interview_active:
                if not self.camera_active or not self.cap:
                    time.sleep(2)
                    continue
                with threading.Lock():  # Add thread safety
                    ret, frame = self.cap.read()
                    if not ret:
                        self._restart_camera()
                        continue
                # Convert to grayscale and apply histogram equalization
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                gray = cv2.equalizeHist(gray)
                # More accurate face detection parameters
                faces = self.face_cascade.detectMultiScale(
                    gray,
                    scaleFactor=1.05,  # Reduced from 1.1
                    minNeighbors=7,    # Increased from 5
                    minSize=(150, 150),# Increased minimum face size
                    flags=cv2.CASCADE_SCALE_IMAGE
                )
                # Only trigger warning if we're very confident
                if len(faces) > 1:
                    # Additional verification - check face sizes are similar
                    areas = [w*h for (x,y,w,h) in faces]
                    if max(areas)/min(areas) < 4:  # Only if faces are similarly sized
                        if not multiple_faces_warning_given:
                            self._handle_cheating_attempt("multiple_faces")
                            multiple_faces_warning_given = True
                    else:
                        multiple_faces_warning_given = False
                else:
                    multiple_faces_warning_given = False
                # Eye detection and attention check
                for (x, y, w, h) in faces:
                    roi_gray = gray[y:y+h, x:x+w]
                    eyes = self.eye_cascade.detectMultiScale(
                        roi_gray,
                        scaleFactor=1.1,
                        minNeighbors=3,
                        minSize=(30, 30))
                    # Only check attention if we have good eye detection
                    if len(eyes) >= 2:  # At least two eyes detected
                        eye_centers = [(ex + ew/2, ey + eh/2) for (ex, ey, ew, eh) in eyes]
                        avg_eye_y = sum(ey for (ex, ey) in eye_centers) / len(eye_centers)
                        # More lenient threshold for looking away
                        if avg_eye_y > h * 0.75 and not looking_away_warning_given:  # Eyes looking down
                            self._handle_cheating_attempt("looking_away")
                            looking_away_warning_given = True
                        elif avg_eye_y <= h * 0.75:
                            looking_away_warning_given = False

    def _restart_camera(self):
        try:
            if self.cap:
                self.cap.release()
            self.cap = cv2.VideoCapture(0)
            self.cap.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
            self.cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)
            self.camera_active = True
        except Exception as e:
            print(f"Camera restart failed: {e}")

    def _monitor_tab_changes(self):
        while not self.tab_monitor_ready:
            time.sleep(0.5)
        try:
            initial_window = gw.getActiveWindow()
            initial_title = initial_window.title if initial_window else "Interview Window"
        except:
            initial_window = None
            initial_title = "Interview Window"
        warning_given = False
        while self.monitoring_active and self.interview_active:
            try:
                current_window = gw.getActiveWindow()
                current_title = current_window.title if current_window else initial_title
                # Only trigger if:
                # 1. We have a valid window
                # 2. The title has actually changed significantly
                # 3. The new title doesn't contain system/keywords
                if (current_window and initial_window and 
                    current_title != initial_title and
                    not any(x in current_title.lower() for x in ["notification", "system", "settings"])):
                    if not warning_given:  # Only warn once per change
                        self.tab_change_detected = True
                        self._handle_cheating_attempt("tab_change")
                        warning_given = True
                else:
                    warning_given = False
                time.sleep(3)  # Longer delay between checks
            except Exception as e:
                print(f"Window monitoring error: {e}")
                time.sleep(3)

    def _handle_cheating_attempt(self, cheat_type):
        """Handle different types of cheating attempts"""
        self.cheating_warnings += 1
        if self.cheating_warnings >= 3:
            self.speak("Multiple concerning behaviors detected. The interview will now conclude.", interruptible=False)
            self.interview_active = False
            return
        responses = {
            "no_face": "Please ensure your face is clearly visible to the camera for the interview.",
            "multiple_faces": "I notice multiple people in the frame. Please ensure you're alone during this interview.",
            "looking_away": "Please maintain focus on the interview and avoid looking at other devices.",
            "tab_change": "Please stay focused on the interview window and avoid switching to other applications."
        }
        if cheat_type in responses:
            self.speak(f"Gentle reminder: {responses[cheat_type]} This is notice {self.cheating_warnings} of 3.", interruptible=False)

    def __del__(self):
        """Clean up resources"""
        self.interview_active = False
        self.monitoring_active = False
        self._stop_camera()
        if hasattr(self, 'face_monitor_thread'):
            self.face_monitor_thread.join(timeout=1)
        if hasattr(self, 'tab_monitor_thread'):
            self.tab_monitor_thread.join(timeout=1)
        pygame.mixer.quit()

    def speak(self, text, interruptible=True):
        if self.interrupted:
            self.interrupted = False
            return
        print(f"Interviewer: {text}")
        try:
            response = self.polly.synthesize_speech(
                Text=text,
                OutputFormat="mp3",
                VoiceId="Aditi"
            )
            if "AudioStream" in response:
                temp_path = os.path.join(tempfile.gettempdir(), f"polly_{int(time.time() * 1000)}.mp3")
                with open(temp_path, 'wb') as f:
                    f.write(response["AudioStream"].read())
                pygame.mixer.music.load(temp_path)
                pygame.mixer.music.play()
                while pygame.mixer.music.get_busy():
                    time.sleep(0.1)
                # Safely attempt file removal
                try:
                    time.sleep(0.2)  # tiny delay before delete
                    os.remove(temp_path)
                except PermissionError:
                    pass  # Suppress WinError 32
        except Exception as e:
            print(f"AWS Polly TTS error: {e}")

    def listen(self, max_attempts=3):
        """Listen for user response with proper context management"""
        for attempt in range(max_attempts):
            try:
                # Create new recognizer instance for this attempt
                attempt_recognizer = sr.Recognizer()
                with self.microphone as source:
                    print("\nListening... (Speak now)")
                    # Adjust for ambient noise with clean context
                    attempt_recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    try:
                        audio = attempt_recognizer.listen(
                            source, 
                            timeout=15, 
                            phrase_time_limit=60
                        )
                        text = attempt_recognizer.recognize_google(audio)
                        print(f"Candidate: {text}")
                        # Process tone detection
                        tone = self._detect_tone(text)
                        if tone != "professional":
                            self.handle_improper_tone(tone)
                            placeholder = "[Response had non-professional tone]"
                            self.conversation_history.append({"role": "user", "content": placeholder})
                            return placeholder
                        if text.strip():
                            self.conversation_history.append({"role": "user", "content": text})
                            return text
                        else:
                            placeholder = "[Unclear response]"
                            self.conversation_history.append({"role": "user", "content": placeholder})
                            return placeholder
                    except sr.WaitTimeoutError:
                        if attempt < max_attempts - 1:
                            self.speak("I didn't hear anything. Please speak when you're ready.", interruptible=False)
                            time.sleep(2)
                        continue
                    except sr.UnknownValueError:
                        if attempt < max_attempts - 1:
                            self.speak("I couldn't quite catch that. Could you please speak again?", interruptible=False)
                            time.sleep(2)
                        continue
                    except sr.RequestError as e:
                        print(f"Speech recognition error: {e}")
                        if attempt < max_attempts - 1:
                            self.speak("There was a technical issue. Please try speaking again.", interruptible=False)
                            time.sleep(2)
                        continue
            except OSError as e:
                print(f"Microphone access error: {e}")
                self.speak("I'm having trouble accessing the microphone. Please check your microphone settings.", interruptible=False)
                placeholder = "[Microphone issue]"
                self.conversation_history.append({"role": "user", "content": placeholder})
                return placeholder
        # If all attempts fail
        placeholder = "[Response unclear after multiple attempts]"
        self.conversation_history.append({"role": "user", "content": placeholder})
        self.speak("Let's continue with the next part of our interview.", interruptible=False)
        return placeholder

    def _rephrase_question(self, question):
        """Rephrase the given question while keeping the same meaning"""
        prompt = f"""Rephrase this interview question to make it clearer while keeping the same meaning:
        Original: {question}
        Requirements:
        - Keep technical accuracy
        - Maintain same difficulty level
        - Don't change the core concept being tested
        - Make it slightly different wording
        - Keep it one sentence
        Return only the rephrased question."""
        rephrased = self.query_gemini(prompt)
        return rephrased.strip() if rephrased else question

    def _detect_tone(self, text):
        if not text:
            return "professional"
        text_lower = re.sub(r'\s+', ' ', text.lower().strip())
        arrogant_keywords = [
            r'\bobviously\b', r'\beveryone knows\b', r'\bchild\'?s play\b',
            r'\bthat\'?s easy\b', r'\btrivial\b', r'\bwaste of time\b'
        ]
        rude_patterns = [
            r'\byou don\'?t understand\b', r'\bthat\'?s stupid\b', r'\bdumb question\b',
            r'\bare you serious\b', r'\bthis is ridiculous\b', r'\bwho cares\b'
        ]
        for pattern in arrogant_keywords:
            if re.search(pattern, text_lower):
                return "arrogant"
        for pattern in rude_patterns:
            if re.search(pattern, text_lower):
                return "rude"
        return "professional"

    def handle_improper_tone(self, tone):
        self.tone_warnings += 1
        if self.tone_warnings >= 2:
            self.speak("I appreciate your participation, but let's maintain a professional tone throughout our conversation.", interruptible=False)
            return
        responses = {
            "arrogant": [
                "I appreciate your confidence! Let's channel that into demonstrating your technical knowledge.",
                "Great confidence! Now let's see how you apply that expertise to solve problems.",
            ],
            "rude": [
                "I understand technical interviews can be stressful. Let's take a moment and continue professionally.",
                "No worries, let's refocus on showcasing your technical abilities.",
            ]
        }
        if tone in responses:
            response = random.choice(responses[tone])
            self.speak(response, interruptible=False)
            time.sleep(1)

    def query_gemini(self, prompt):
        try:
            response = self.model.generate_content(prompt)
            if hasattr(response, 'text'):
                return response.text
            elif hasattr(response, 'result'):
                return response.result
            elif hasattr(response, 'candidates') and response.candidates:
                return response.candidates[0].content.parts[0].text
            else:
                return "Could you tell me more about your experience with that?"
        except Exception as e:
            print(f"Gemini API Error: {e}")
            return "Could you elaborate on your experience with that technology?"

    def _identify_tech_domain(self, text):
        if not text:
            return None
        text_lower = text.lower()
        domain_scores = {domain: 0 for domain in self.tech_domains}
        for domain, keywords in self.tech_domains.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    domain_scores[domain] += 2
                elif re.search(r'\b' + re.escape(keyword.lower()) + r'\b', text_lower):
                    domain_scores[domain] += 1
        best_tech_domain, tech_score = max(domain_scores.items(), key=lambda x: x[1])
        if tech_score >= 2:
            print(f"[Domain Detection] Detected technical domain: {best_tech_domain}")
            return best_tech_domain
        # Check non-technical domains
        domain_scores = {domain: 0 for domain in self.non_tech_domains}
        for domain, keywords in self.non_tech_domains.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    domain_scores[domain] += 2
                elif re.search(r'\b' + re.escape(keyword.lower()) + r'\b', text_lower):
                    domain_scores[domain] += 1
        best_non_tech_domain, non_tech_score = max(domain_scores.items(), key=lambda x: x[1])
        if non_tech_score >= 2:
            print(f"[Domain Detection] Detected non-technical domain: {best_non_tech_domain}")
            return best_non_tech_domain
        # If ambiguous but tech keywords exist, fallback
        fallback_keywords = ["pandas", "numpy", "machine learning", "tensorflow", "keras", "python", "scikit", "data pipeline", "model", "deployment"]
        if any(kw in text_lower for kw in fallback_keywords):
            print("[Domain Detection] Fallback to technical domain: data science")
            return "data science"
        print("[Domain Detection] No strong domain match found.")
        return None

    def _generate_coding_question(self, domain, difficulty="medium"):
        """Generate a coding question based on the candidate's domain"""
        domain_mapping = {
            "python": "Python",
            "java": "Java", 
            "cpp": "C++",
            "frontend": "JavaScript",
            "backend": "Python or your preferred language",
            "AI": "Python",
            "data science": "Python",
            "machine learning": "Python"
        }
        language = domain_mapping.get(domain, "Python")
        prompt = f"""Generate a {difficulty} level coding problem suitable for a technical interview in {domain}.
        Requirements:
        - Should be solvable in {language}
        - Should take 10-15 minutes to solve
        - Include a clear problem statement
        - Provide input/output examples
        - Should test algorithmic thinking and {domain} knowledge
        - Avoid problems that are too easy or too hard
        - Focus on practical problem-solving skills
        Format your response as:
        Problem: [Clear problem statement]
        Example Input: [Sample input]
        Example Output: [Expected output]
        Constraints: [Any constraints or edge cases to consider]
        Generate only the problem, no solution."""
        try:
            response = self.query_gemini(prompt)
            return response.strip() if response else None
        except Exception as e:
            print(f"Error generating coding question: {e}")
            return self._get_fallback_coding_question(domain)

    def _get_fallback_coding_question(self, domain):
        """Fallback coding questions if AI generation fails"""
        fallback_questions = {
            "python": """Problem: Find the two numbers in a list that add up to a target sum.
Example Input: numbers = [2, 7, 11, 15], target = 9
Example Output: [0, 1] (indices of numbers 2 and 7)
Constraints: Each input has exactly one solution, and you may not use the same element twice.""",
            "default": """Problem: Write a function to reverse words in a sentence while keeping the word order.
Example Input: "Hello World Python"
Example Output: "olleH dlroW nohtyP"
Constraints: Preserve spaces between words, handle empty strings gracefully."""
        }
        return fallback_questions.get(domain, fallback_questions["default"])

    def _coding_followup(self, code, language):
        """Ask follow-up questions about the code submitted by the candidate."""
        prompt = f"""You are an expert software engineer reviewing code written in {language}.
        The candidate has provided the following code:
        ```
{code}
    Ask one follow-up question that:
    - Tests their understanding of the code they wrote
    - Explores potential edge cases or improvements
    - Is specific to the code provided
    - Is clear and concise
    - Can be answered without running the code
    - Focuses on code clarity, efficiency, or potential bugs
    - Is appropriate for an interview setting
    Generate only the question, no additional text."""
        try:
            response = self.query_gemini(prompt)
            return response.strip() if response else None
        except Exception as e:
            print(f"Error generating coding follow-up: {e}")
            return "Can you walk me through your code and explain your approach?"

    def _save_transcription_to_docx(self, file_path="interview_transcript.docx"):
        doc = Document()
        doc.add_heading("Interview Transcription", level=1)
        for i, msg in enumerate(self.conversation_history):
            role = msg.get("role", "unknown").capitalize()
            content = msg.get("content", "")
            doc.add_paragraph(f"{role}: {content}")
        doc.save(file_path)
        print(f"Transcript saved to {file_path}")
        return file_path

    def _generate_feedback_from_docx(self, docx_path="interview_transcript.docx"):
        try:
                doc = Document(docx_path)
                transcript_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                prompt = f"""
                Below is a transcription of an interview. Perform the following tasks STRICTLY as described:
                1. Extract the interviewer's questions and the candidate's answers:
                - Ignore filler words like "okay," "hmm," "uh," etc., unless part of a meaningful question/answer.
                - Only include complete sentences or meaningful phrases.
                - Skip any exchange where:
                - The question is filler (e.g., "Okay", "Hmm", "Got it").
                - The answer is too short, incomplete, or irrelevant (e.g., "Yes", "No", "Maybe", "I think so").
                2. Categorize each question under a generalized skill category (e.g., Python, AI, JavaScript, Machine Learning, etc.).
                3. For each skill category:
                - Summarize the candidate's performance concisely (STRICTLY within 900 characters). If exceeded, truncate and add "..." at the end.
                4. For each extracted question-answer pair:
                - Include start and end timestamps in seconds (relative to interview start).
                - Group questions by skill area into a single block.
                - Ensure "que" (question) is within 900 characters. If exceeded, truncate and add "...".
                - Ensure "ans" (answer) is within 4000 characters. If exceeded, truncate and add "...".
                - Also provide a concise alternate ideal answer (under 1000 characters) if the answer is weak.
                5. Provide an overall evaluation:
                - Candidate strengths (STRICTLY within 400 characters). If exceeded, truncate and add "...".
                - Points of improvement (STRICTLY within 400 characters). If exceeded, truncate and add "..."
                6. Rate the candidate on:
                - Communication: Choose one — poor, average, good, excellent.
                - Attitude: Choose one — poor, average, good, excellent.
                7. If the candidate asked to repeat a question or said they didn't understand:
                - Indicate that and include the **rephrased question** that was asked instead.
                - If they still could not answer, generate a **model answer** that could have helped.
                - Mark such entries under `"alternate_explanation": "..."` in the output.
                8. Create a feedback form titled `"feedback_form"` containing:
                - `"summary"`: a concise paragraph (300–500 characters) summarizing the overall performance
                - `"key_areas_to_improve"`: list of 3 suggestions
                - `"recommended_study_links"`: list of URLs or topics (max 5)
                ...
                Transcription:
                {transcript_text}
                """
                feedback = self.query_gemini(prompt)
                with open("final_interview_feedback.json", "w") as f:
                    f.write(feedback.strip())
                    print("Feedback saved to final_interview_feedback.json")
        except Exception as e:
            print(f"[ERROR] Feedback generation failed: {e}")

    def start_interview(self):
        # Start tab monitoring
        def enable_tab_monitor():
            time.sleep(3)
            self.tab_monitor_ready = True
        threading.Thread(target=enable_tab_monitor, daemon=True).start()
        # Start the interview logic
        interview_thread = threading.Thread(target=self._run_interview_logic)
        interview_thread.daemon = True
        interview_thread.start()
        # Keep the main thread alive while interview is active
        while self.interview_active:
            time.sleep(1)
class RAGExpertTechnicalInterviewer(ExpertTechnicalInterviewer):
        def __init__(self, model="gemini-2.0-flash", accent="indian"):
            super().init (model, accent)
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2') # Pre-trained embedding model
            self.knowledge_base_path = "knowledge_base.json"
            self.vector_index_path = "vector_index.faiss"
            self.vector_dimension = 384 # MiniLM-L6 outputs 384-dimensional vectors
            self.vector_index = self._load_or_create_vector_index()
            self.knowledge_base = self._load_knowledge_base()

        def _load_or_create_vector_index(self):
            """Load an existing FAISS index or create a new one."""
            if os.path.exists(self.vector_index_path):
                return faiss.read_index(self.vector_index_path)
            else:
                return faiss.IndexFlatL2(self.vector_dimension)

        def _load_knowledge_base(self):
            """Load the knowledge base from a JSON file."""
            if os.path.exists(self.knowledge_base_path):
                with open(self.knowledge_base_path, "r") as f:
                    return json.load(f)
            return []

        def _save_knowledge_base(self):
            """Save the knowledge base to a JSON file."""
            with open(self.knowledge_base_path, "w") as f:
                json.dump(self.knowledge_base, f, indent=4)

        def _save_vector_index(self):
            """Save the FAISS index to disk."""
            faiss.write_index(self.vector_index, self.vector_index_path)

        def _add_to_knowledge_base(self, text):
            """Add new text to the knowledge base and update the vector index."""
            embedding = self.embedding_model.encode(text)
            self.knowledge_base.append({"text": text, "embedding": embedding.tolist()})
            self.vector_index.add(np.array([embedding]))
            self._save_knowledge_base()
            self._save_vector_index()

        def _retrieve_context(self, query, top_k=3):
            """Retrieve the top-k most relevant documents from the knowledge base."""
            query_embedding = self.embedding_model.encode(query)
            distances, indices = self.vector_index.search(np.array([query_embedding]), top_k)
            retrieved_texts = [self.knowledge_base[i]["text"] for i in indices[0]]
            return retrieved_texts

        def query_gemini_with_rag(self, prompt, query):
            """Query the generative model with additional context from the knowledge base."""
            retrieved_context = self._retrieve_context(query)
            full_prompt = f"{prompt}\nAdditional Context:\n" + "\n".join(retrieved_context)
            return self.query_gemini(full_prompt)

        def _update_knowledge_base_after_interview(self):
            """Update the knowledge base with the latest conversation history."""
            for msg in self.conversation_history:
                self._add_to_knowledge_base(msg["content"])

        def _run_interview_logic(self):
            try:
                super()._run_interview_logic()
            finally:
        # Update the knowledge base after the interview ends
                self._update_knowledge_base_after_interview()
                docx_path = self._save_transcription_to_docx()
                self._generate_feedback_from_docx(docx_path)
        if __name__ == "__main__":
            try:
                interviewer = ExpertTechnicalInterviewer()
                if os.path.exists("client_questions.json"):
                    interviewer.load_client_questions("client_questions.json")
                interviewer.start_interview()
            except Exception as e:
                print(f"Fatal error: {e}")
        print(f"Fatal error: {e}")
        print("Please check your environment setup:")
        print("1. GEMINI_API_KEY in .env file")
        print("2. Microphone and camera permissions")
        print("3. Required Python packages installed")
        print("\nInterview session ended. Camera automatically turned off.")
        print("Thank you for using the Enhanced Technical Interview Bot!")